"""Knowledge base document API."""

import io
from uuid import UUID

from docx import Document as DocxDocument
from fastapi import APIRouter, File, HTTPException, UploadFile
from pypdf import PdfReader

from app.core.config import get_settings
from app.dependencies import get_document_ingestor
from app.models.document import Document, DocumentListResponse, IngestResponse, UrlIngestRequest
from app.db.database import get_supabase_client

router = APIRouter(prefix="/documents", tags=["documents"])


def _org_id() -> str:
    settings = get_settings()
    if settings.default_organization_id:
        return settings.default_organization_id
    db = get_supabase_client()
    org = db.table("organizations").select("id").limit(1).execute()
    if not org.data:
        raise HTTPException(status_code=500, detail="No organization configured")
    return org.data[0]["id"]


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


@router.get("", response_model=DocumentListResponse)
async def list_documents():
    db = get_supabase_client()
    org_id = _org_id()
    result = (
        db.table("documents")
        .select("*")
        .eq("organization_id", org_id)
        .order("created_at", desc=True)
        .execute()
    )
    items = [Document(**d) for d in (result.data or [])]
    return DocumentListResponse(items=items, total=len(items))


@router.post("/upload", response_model=IngestResponse)
async def upload_document(file: UploadFile = File(...)):
    filename = file.filename or "document"
    content_bytes = await file.read()
    lower = filename.lower()

    if lower.endswith(".pdf"):
        text = _extract_pdf(content_bytes)
        source_type = "pdf"
    elif lower.endswith((".doc", ".docx")):
        text = _extract_docx(content_bytes)
        source_type = "doc"
    else:
        raise HTTPException(status_code=400, detail="Supported formats: PDF, DOC, DOCX")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    ingestor = get_document_ingestor()
    from uuid import UUID as PyUUID

    doc_id, chunks = ingestor.ingest(
        organization_id=PyUUID(_org_id()),
        title=filename,
        content=text,
        source_type=source_type,
    )
    return IngestResponse(
        document_id=doc_id,
        chunks_created=chunks,
        message=f"Ingested {chunks} chunks from {filename}",
    )


@router.post("/url", response_model=IngestResponse)
async def ingest_url(body: UrlIngestRequest):
    import httpx
    from bs4 import BeautifulSoup

    url = str(body.url)
    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
        response = await client.get(url)
        response.raise_for_status()

    soup = BeautifulSoup(response.text, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    if not text:
        raise HTTPException(status_code=400, detail="No text extracted from URL")

    title = body.title or url
    ingestor = get_document_ingestor()
    from uuid import UUID as PyUUID

    doc_id, chunks = ingestor.ingest(
        organization_id=PyUUID(_org_id()),
        title=title,
        content=text,
        source_type="url",
        source_url=url,
    )
    return IngestResponse(
        document_id=doc_id,
        chunks_created=chunks,
        message=f"Ingested {chunks} chunks from URL",
    )


@router.delete("/{document_id}")
async def delete_document(document_id: UUID):
    db = get_supabase_client()
    db.table("document_chunks").delete().eq("document_id", str(document_id)).execute()
    db.table("documents").delete().eq("id", str(document_id)).execute()
    return {"status": "deleted"}
